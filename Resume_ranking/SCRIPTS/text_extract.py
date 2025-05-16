import os
import fitz
from docx import Document
# from pdfminer.high_level import extract_text


def check_file_type(file):
    # Extract text based on file type
    filename = os.path.basename(file)
    if file.lower().endswith(".pdf"):
        return ".pdf"
    elif file.lower().endswith(".docx"):
        return ".docx"
    else:
        print(f"Skipping {file}: Unsupported file format!")




def extract_from_pdf(pdf_file):
    text = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text


def extract_from_docx(docx_file):
    text = ""
    document = Document(docx_file)

    # Extract paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"

    # Extract text inside tables (resumes may have tables for skills, etc.)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text.strip() + "\n"

    return text

import re

def preprocess_text(text):
    """Enhanced preprocessing for resume text"""

    # Remove control/non-printable characters
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)

    # Normalize commas and semicolons (space padded)
    text = re.sub(r'[,\;]', ' , ', text)

    # Normalize various bullets/symbols
    text = re.sub(r'[•\-\*\+‣⁃▪■●♦]', ' • ', text)

    # Replace multiple spaces with single space (but preserve line breaks)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n+', '\n', text)

    # Remove redundant symbols (e.g., ====, ----, etc.)
    text = re.sub(r'[-=]{2,}', '', text)

    # Lowercase everything
    text = text.lower()

    # Final strip
    return text.strip()




